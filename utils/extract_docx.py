from docx import Document
import logging

logger = logging.getLogger(__name__)

def extract_text(file_path):
    try:
        logger.info(f"Extracting text from Word document: {file_path}")
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Only add non-empty cells
                        text += cell.text + "\n"
        
        if not text.strip():
            logger.warning(f"No text content found in document: {file_path}")
            return "No text content found in document."
            
        logger.info(f"Successfully extracted text from Word document: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from Word document {file_path}: {str(e)}")
        raise 